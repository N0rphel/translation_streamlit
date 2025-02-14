import os
import re
import pandas as pd
import streamlit as st
from io import BytesIO
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM

# PDF and DOCX processing
import PyPDF2
import docx

# Constants
MODELS_DIR = "./models"

# Utility functions
def process_txt(text):
    paragraphs = re.split(r'\n\s*\n+', text.strip())
    return [para.strip() for para in paragraphs if para.strip()]

def translate_paragraph(text, model, tokenizer):
    paragraphs = process_txt(text)
    translated_paragraphs = []
    
    for para in paragraphs:
        sentences = re.split(r'(?<=[.!?])\s+', para.strip())
        translated_sentences = [translate(sentence, model, tokenizer) for sentence in sentences]
        translated_paragraphs.append(' '.join(translated_sentences))
    
    return '\n\n'.join(translated_paragraphs)

def list_models():
    return [f for f in os.listdir(MODELS_DIR) if os.path.isdir(os.path.join(MODELS_DIR, f))] if os.path.exists(MODELS_DIR) else []

@st.cache_resource
def load_model(path):
    tokenizer = AutoTokenizer.from_pretrained(path)
    model = AutoModelForSeq2SeqLM.from_pretrained(path)
    return model, tokenizer

def translate(text, model, tokenizer):
    inputs = tokenizer(text, return_tensors="pt").input_ids
    translated_tokens = model.generate(inputs, max_new_tokens=512)
    return tokenizer.batch_decode(translated_tokens, skip_special_tokens=True)[0]

# Streamlit UI
st.title("Translation App")

# Load Model Section
with st.expander("Load Model"):
    st.write("The models should be placed in the models directory!!!")
    models = list_models()
    
    if models:
        selected_model = st.selectbox("Available Models:", models)
        model_path = os.path.join(MODELS_DIR, selected_model).replace("\\", "/")
        st.write(f"**Selected Model:** {selected_model}")
        
        if st.button("Load Model"):
            model, tokenizer = load_model(model_path)
            st.session_state["model"], st.session_state["tokenizer"] = model, tokenizer
            st.success(f"✅ Model '{selected_model}' loaded successfully!")
    else:
        st.warning("No models found in the `/models/` directory.")

# Translate Section
with st.expander("Translate"):
    text_to_translate = st.text_area("Enter text to translate")
    if st.button("Translate") and text_to_translate:
        if "model" in st.session_state and "tokenizer" in st.session_state:
            translate_func = translate_paragraph if "\n" in text_to_translate or len(text_to_translate.split('.')) > 1 else translate
            translated_text = translate_func(text_to_translate, st.session_state["model"], st.session_state["tokenizer"])
            st.text_area("Translated Text", value=translated_text, height=200)
        else:
            st.error("Please load a model first.")

# Batch Translation Section
with st.expander("Batch Translation"):
    st.write("Accepted file types: CSV, TXT, PDF, DOCX")
    translation_option = st.radio("Select Translation Mode", ["Line-by-Line Translation", "Paragraph Translation"])
    uploaded_file = st.file_uploader("Choose a file", type=['csv', 'txt', 'pdf', 'docx'])

    if uploaded_file:
        file_ext = uploaded_file.name.split('.')[-1]
        text_data = ""
        
        if file_ext == "txt":
            text_data = uploaded_file.getvalue().decode("utf-8")
        elif file_ext == "pdf":
            text_data = "".join(page.extract_text() + "\n" for page in PyPDF2.PdfReader(uploaded_file).pages)
        elif file_ext == "docx":
            text_data = "\n".join(para.text for para in docx.Document(uploaded_file).paragraphs)
        
        if file_ext == "csv":
            df = pd.read_csv(uploaded_file, header=None, names=['Source'])
            st.write(df)
            
            if st.button("Batch Translate CSV") and "model" in st.session_state:
                df["Translation"] = df["Source"].apply(lambda x: translate(str(x), st.session_state["model"], st.session_state["tokenizer"])) if translation_option == "Line-by-Line Translation" else pd.DataFrame({"Source": [" ".join(df["Source"].tolist())], "Translation": [translate_paragraph(" ".join(df["Source"].tolist()), st.session_state["model"], st.session_state["tokenizer"]) ]})
                st.write("Translated Data:", df.head())
                st.download_button("Download Translated CSV", df.to_csv(index=False).encode('utf-8'), "Translated.csv", "text/csv")
            else:
                st.error("Please load a model first.")
        else:
            if translation_option == "Line-by-Line Translation":
                sentences = process_txt(text_data)
                df = pd.DataFrame(sentences, columns=['Source'])
                st.write(df)
                
                if st.button("Batch Translate Document") and "model" in st.session_state:
                    df["Translation"] = df["Source"].apply(lambda x: translate(str(x), st.session_state["model"], st.session_state["tokenizer"]))
                    st.download_button("Download Translated CSV", df.to_csv(index=False).encode('utf-8'), "Translated.csv", "text/csv")
                else:
                    st.error("Please load a model first.")
            else:
                if st.button("Batch Translate Document (Paragraph)") and "model" in st.session_state:
                    translation = translate_paragraph(text_data, st.session_state["model"], st.session_state["tokenizer"])
                    st.write("**Translated Text:**", translation)
                    filename_base = f"Translated_{uploaded_file.name.split('.')[0]}"
                    
                    st.download_button("Download as TXT", translation, f"{filename_base}.txt", "text/plain")
                    doc_io = BytesIO()
                    docx.Document().add_paragraph(translation).save(doc_io)
                    doc_io.seek(0)
                    st.download_button("Download as DOCX", doc_io.getvalue(), f"{filename_base}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                else:
                    st.error("Please load a model first.")

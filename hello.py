import streamlit as st
import os
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM
import pandas as pd
import re


# Import libraries for PDF and DOCX processing
import PyPDF2
import docx
from io import BytesIO

MODELS_DIR = "./models" 


def process_txt(text):
    # Split by one or more newlines (with optional whitespace) to capture paragraphs
    paragraphs = re.split(r'\n\s*\n+', text.strip())
    return [para.strip() for para in paragraphs if para.strip()]

def translate_paragraph(text, model, tokenizer):
    paragraphs = process_txt(text)  # Split input into paragraphs
    translated_paragraphs = []
    
    for para in paragraphs:
        # Split paragraph into sentences
        sentences = re.split(r'(?<=[.!?])\s+', para.strip())
        translated_sentences = [translate(sentence, model, tokenizer) for sentence in sentences]
        
        # Join sentences into a paragraph
        translated_paragraph = ' '.join(translated_sentences)
        translated_paragraphs.append(translated_paragraph)
    
    # Join paragraphs with double newlines to preserve spacing
    return '\n\n'.join(translated_paragraphs)

# For listing all the models in the models folder
def list_models():
    if os.path.exists(MODELS_DIR):
        return [f for f in os.listdir(MODELS_DIR) if os.path.isdir(os.path.join(MODELS_DIR, f))]
    return []

@st.cache_resource
def load_model(path):
    st.write(path)
    tokenizer = AutoTokenizer.from_pretrained(path,)
    model = AutoModelForSeq2SeqLM.from_pretrained(path)
    return model, tokenizer

def translate(text, model, tokenizer):
    inputs = tokenizer(text, return_tensors="pt").input_ids
    translated_tokens = model.generate(inputs, max_new_tokens=512)
    return tokenizer.batch_decode(translated_tokens, skip_special_tokens=True)[0]

st.title("Translation App")

with st.expander("Load Model:"):
    st.write("The models should be placed in the models directory!!!")
    st.write("Select a model from the list below:")
    
    models = list_models()
    if models:
        selected_model = st.selectbox("Available Models:", models)
        model_path = os.path.join(MODELS_DIR, selected_model).replace("\\", "/")
        st.write(f"**Selected Model:** {selected_model}")
        st.write(f"**Model Path:** {model_path}")

        if st.button("Load Model"):
            model, tokenizer = load_model(model_path)
            st.session_state["model"] = model
            st.session_state["tokenizer"] = tokenizer
            st.success(f"✅ Model '{selected_model}' loaded successfully!")
    else:
        st.write("No models found in the `/models/` directory.")

with st.expander("Translate:"):
    text_to_translate = st.text_area("Enter text to translate")
    if st.button("Translate"):
        if text_to_translate:
            if "model" in st.session_state and "tokenizer" in st.session_state:
                if "\n" in text_to_translate or len(text_to_translate.split('.')) > 1:
                    translated_text = translate_paragraph(text_to_translate,
                                                           st.session_state["model"],
                                                           st.session_state["tokenizer"])
                else:
                    translated_text = translate(text_to_translate,
                                                st.session_state["model"],
                                                st.session_state["tokenizer"])
                st.text_area("Translated Text", value=translated_text, height=200)
            else:
                st.error("Please load a model first.")
        else:
            st.warning("Please enter text to translate.")

with st.expander("Batch Translation:"):
    st.write("Accepted file types: CSV, TXT, PDF, DOCX")
    st.write("• CSV: Expected to have rows of sentences\n• TXT: Contains paragraphs or sentences\n• PDF/DOCX: The text will be extracted from the document")

    # Let the user select the translation mode
    translation_option = st.radio("Select Translation Mode", 
                                  options=["Line-by-Line Translation", "Paragraph Translation"])

    # Display a simple diagram to illustrate the mode
    if translation_option == "Line-by-Line Translation":
        st.markdown("**Diagram: Line-by-Line Translation**")
        st.markdown("""
        ```
        [Line 1] -> [Translated Line 1]
        [Line 2] -> [Translated Line 2]
        [Line 3] -> [Translated Line 3]
        ```
        """)
    else:
        st.markdown("**Diagram: Paragraph Translation**")
        st.markdown("""
        ```
        [Entire Paragraph] -----> [Split into sentences]
                 ^                           |
                 |                           v
        [Recombine into paragraph] <---- [Translate each sentence]
        ```
        """)

    uploaded_file = st.file_uploader("Choose a file", type=['csv', 'txt', 'pdf', 'docx'])

    if uploaded_file is not None:
        # --- Process CSV Files ---
        if uploaded_file.name.endswith(".csv"):
            df = pd.read_csv(uploaded_file, header=None, names=['Source'])
            st.write(df)

            if st.button("Batch Translate CSV"):
                if "model" in st.session_state and "tokenizer" in st.session_state:
                    if translation_option == "Line-by-Line Translation":
                        df["Translation"] = df["Source"].apply(
                            lambda x: translate(str(x),
                                                st.session_state["model"],
                                                st.session_state["tokenizer"])
                        )
                    else:
                        # For paragraph translation, join all rows into a single paragraph
                        joined_text = " ".join(df["Source"].tolist())
                        translation = translate_paragraph(joined_text,
                                                          st.session_state["model"],
                                                          st.session_state["tokenizer"])
                        df = pd.DataFrame({"Source": [joined_text],
                                           "Translation": [translation]})
                    st.write("Translated Data:", df.head())
                    csv_data = df.to_csv(index=False, header=True).encode('utf-8')
                    st.download_button("Download Translated CSV", csv_data, "Translated.csv", "text/csv")
                else:
                    st.error("Please load a model first.")

        # --- Process TXT, PDF, and DOCX Files ---
        elif uploaded_file.name.endswith(".txt") or \
             uploaded_file.name.endswith(".pdf") or \
             uploaded_file.name.endswith(".docx"):
            
            # Extract text based on file type
            if uploaded_file.name.endswith(".txt"):
                text_data = uploaded_file.getvalue().decode("utf-8")
            elif uploaded_file.name.endswith(".pdf"):
                pdf_reader = PyPDF2.PdfReader(uploaded_file)
                text_data = ""
                for page in pdf_reader.pages:
                    text_data += page.extract_text() + "\n"
            elif uploaded_file.name.endswith(".docx"):
                doc = docx.Document(uploaded_file)
                text_data = "\n".join([para.text for para in doc.paragraphs])
            
            # Process based on translation option
            if translation_option == "Line-by-Line Translation":
                sentences = process_txt(text_data)
                df = pd.DataFrame(sentences, columns=['Source'])
                st.write(df)
                if st.button("Batch Translate Document (Line-by-Line)"):
                    if "model" in st.session_state and "tokenizer" in st.session_state:
                        df["Translation"] = df["Source"].apply(
                            lambda x: translate(str(x),
                                                st.session_state["model"],
                                                st.session_state["tokenizer"])
                        )
                        st.write("Translated Data:", df.head())
                        csv_data = df.to_csv(index=False, header=True).encode('utf-8')
                        st.download_button("Download Translated CSV", csv_data, "Translated.csv", "text/csv")
                    else:
                        st.error("Please load a model first.")
            else:
                # Batch translation (after translating)
                if translation_option == "Paragraph Translation":
                    # For paragraph translation, treat the entire document as one block
                    st.write("**Original Document Text:**")
                    st.write(text_data)

                    if st.button("Batch Translate Document (Paragraph)"):
                        if "model" in st.session_state and "tokenizer" in st.session_state:
                            translation = translate_paragraph(
                                text_data,
                                st.session_state["model"],
                                st.session_state["tokenizer"]
                            )
                            
                            # Store the translation in session state
                            st.session_state["translation"] = translation
                            
                            st.write("**Translated Text:**")
                            st.write(translation)

                            base_filename = os.path.splitext(uploaded_file.name)[0]
                            new_filename = f"Translated_{base_filename}"

                            # ✅ Download as TXT
                            st.download_button(
                                "Download as TXT",
                                translation,
                                new_filename + ".txt",
                                "text/plain"
                            )

                            # ✅ Generate DOCX using python-docx
                            doc = docx.Document()
                            doc.add_paragraph(translation)

                            docx_io = BytesIO()
                            doc.save(docx_io)
                            docx_io.seek(0)

                            st.download_button(
                                "Download as DOCX",
                                docx_io.getvalue(),
                                new_filename + ".docx",
                                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )

                        else:
                            st.error("Please load a model first.")
